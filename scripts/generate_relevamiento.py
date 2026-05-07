"""Genera `Relevamiento-Inicial.docx` — el punto de partida del proyecto:
contexto del cliente, sistema viejo, pain points, requerimientos funcionales
del PMV (RF-01..RF-30), decisiones arquitectónicas y roadmap de fases.

Uso:
    cd backend
    .venv/Scripts/python ../../scripts/generate_relevamiento.py
"""
from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "manuals" / "output" / "Relevamiento-Inicial.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_rf_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    """rows = [(rf_id, modulo, descripcion)]."""
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    for i, h in enumerate(("ID", "Módulo", "Requerimiento")):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "1F2937")

    for i, (rf_id, modulo, desc) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = rf_id
        cells[1].text = modulo
        cells[2].text = desc
        for r in cells[0].paragraphs[0].runs:
            r.bold = True


def main() -> int:
    doc = Document()
    today = dt.date.today().isoformat()

    # ================== Portada ==================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Relevamiento Inicial")
    run.font.size = Pt(32)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema de Gestión Integral")
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CASA SALCO — Río Cuarto, Córdoba")
    run.font.size = Pt(16)
    run.italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Fecha de relevamiento: {today}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento de inicio — del relevamiento nacen los requerimientos")
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ================== Resumen ejecutivo ==================
    add_heading(doc, "1. Resumen ejecutivo", level=1)
    add_paragraph(
        doc,
        "CASA SALCO es un comercio mayorista/minorista multi-sucursal de Río "
        "Cuarto, Córdoba, dedicado al rubro almacén con áreas de fiambrería, "
        "pollería, drugstore, kiosco/cigarrillos, cuidado personal, "
        "descartables y otros. Operativamente trabaja con 4 sucursales "
        "(originalmente Centro, Norte, Sur, Express) y un catálogo de unos "
        "35.000 artículos provistos por ~140 proveedores.",
    )
    add_paragraph(
        doc,
        "El sistema actual del cliente es una aplicación legacy escrita en "
        "Harbour (xBase / Clipper moderno) con base de datos DBF, que viene "
        "operando desde hace años. Esa solución ya no escala con la realidad "
        "del comercio: integraciones modernas (Mercado Pago, AFIP, OCR), "
        "operación remota y reporting eficiente quedaron afuera. El cliente "
        "está dispuesto a invertir en una solución moderna con la condición "
        "operativa de que le resuelva trabajo administrativo concreto — "
        "su frase textual es: \"prefiero pagar más por ahorrarme un empleado "
        "administrativo\".",
    )
    add_paragraph(
        doc,
        "Del relevamiento surgieron 30 requerimientos funcionales (RF-01 a "
        "RF-30) agrupados en 9 módulos, un dataset legacy real (dos archivos "
        ".xls con ~78.000 filas) que sirve de fuente de verdad para iniciar "
        "el sistema, y un roadmap de 4 fases progresivas: MVP de sincronización "
        "y catálogo, POS y AFIP, dashboard premium, y finalmente OCR/datamining.",
    )

    doc.add_page_break()

    # ================== 2. Contexto y cliente ==================
    add_heading(doc, "2. Contexto y antecedentes", level=1)

    add_heading(doc, "2.1 El cliente", level=2)
    add_bullets(
        doc,
        [
            "Razón social: CASA SALCO (Río Cuarto, Córdoba, Argentina).",
            "Operación: comercio mayorista/minorista del rubro almacén.",
            "Sucursales activas: 4 — Centro, Norte, Sur y Express. La operación se centraliza administrativamente y cada sucursal funciona como punto de venta.",
            "Catálogo: ~35.000 artículos, ~140 proveedores, ~610 rubros y ~2.050 marcas distintas en data legacy.",
            "Áreas internas por sucursal: Comestibles (almacén), Drugstore, Fiambrería y Pollería (más cigarrillos, kiosco, cotillón, pinturería, cuidado personal, descartables, etc. provenientes de la categorización legacy).",
        ],
    )

    add_heading(doc, "2.2 Driver del cliente", level=2)
    add_paragraph(
        doc,
        'Frase textual del cliente: "prefiero pagar más por ahorrarme un '
        'empleado administrativo". Esta declaración prioriza features que '
        "automaticen tareas administrativas concretas (compras, conciliación, "
        "OCR, integraciones de cobro), por encima de mejoras puramente "
        "estéticas o de pulido de pantalla.",
        italic=True,
    )

    add_heading(doc, "2.3 Stakeholders y roles operativos", level=2)
    add_bullets(
        doc,
        [
            "Dueño / Administrador: define configuración del sistema, ABM de usuarios, parámetros del negocio, lectura de reportes.",
            "Supervisor: gestiona operación diaria — artículos, proveedores, stock, compras.",
            "Cajero: opera el Punto de Venta en mostrador.",
            "Fiambrero / Repositor: roles operativos especializados (fiambrería con balanza, reposición de góndola).",
            "Contador: trabaja con información fiscal y financiera (facturas, IVA, cuentas corrientes, exportaciones).",
        ],
    )

    doc.add_page_break()

    # ================== 3. Sistema actual ==================
    add_heading(doc, "3. Sistema actual (legacy)", level=1)

    add_heading(doc, "3.1 Tecnología", level=2)
    add_bullets(
        doc,
        [
            "Lenguaje: Harbour (xBase / Clipper moderno).",
            "Base de datos: archivos DBF independientes por entidad.",
            "UI: terminal estilo años 90 con teclas de función (F1, F2, F3...).",
            "Distribución: instalación local en cada sucursal sin sync nativo entre puntos.",
        ],
    )

    add_heading(doc, "3.2 Convenciones del usuario heredadas del viejo", level=2)
    add_bullets(
        doc,
        [
            "F3 = consultas. La pantalla \"Consultas\" del nuevo sistema mantiene esa metáfora.",
            "Mantenimiento = artículos / catálogo / configuración.",
            "Cobranzas = cuenta corriente + cheques.",
            "Listados = reportes.",
            "Caja = movimientos de caja del día.",
        ],
    )

    add_heading(doc, "3.3 Antipatrones detectados (Dragonfish-like)", level=2)
    add_paragraph(
        doc,
        "Algunas convenciones del sistema viejo se identificaron como puntos "
        "de mejora explícitos del nuevo sistema. El más recurrente:",
    )
    add_bullets(
        doc,
        [
            "Vista de compras NO agrupada por proveedor: en el viejo, listado plano de movimientos sin distinción clara por proveedor. En el nuevo, la vista de compras se agrupará/filtrará por proveedor con búsqueda inline.",
            "Productos identificados con un solo código de barras: el viejo no soporta múltiples códigos (principal + alternos por proveedor + empaquetados). En el nuevo, tabla 1:N de códigos.",
            "Datos sucios en las tablas: filas de prueba, productos con código=nombre, registros inconsistentes. El nuevo sistema valida e importa con reporte de saltadas.",
        ],
    )

    add_heading(doc, "3.4 Datos legacy disponibles", level=2)
    add_paragraph(
        doc,
        "El cliente entregó dos archivos Excel exportados del sistema viejo, "
        "que constituyen la fuente de verdad CURRENT a importar al nuevo "
        "sistema. Son la base con la que arranca el ETL.",
    )
    add_bullets(
        doc,
        [
            "Archivo 1 (proveedores y relación articulo-proveedor): 22.683 filas en hoja EMPAQUETADOS · 43.433 filas en hoja RELACION ARTICULO-PROVEEDOR · 144 filas en hoja PROVEEDOR.",
            "Archivo 2 (maestro de artículos): 34.843 artículos con descripción, costo, PVP, rubro/grupo, marca, proveedor principal.",
            "Top proveedores por volumen: MAXICONSUMO (~4.200 artículos), DIARCO (~4.180), D+D (~2.330), DISTRI TDF (~2.330).",
            "Encoding mal codificado (BIFF cp1252) — requiere re-decode en Python al importar.",
            "Datos sucios identificados: filas de prueba (\"ASDLASJDLAS\"), código=nombre (proveedor \"muerto\"), 798 artículos con costo=0.",
        ],
    )

    doc.add_page_break()

    # ================== 4. Pain points ==================
    add_heading(doc, "4. Pain points y oportunidad", level=1)

    add_paragraph(
        doc,
        "Problemas concretos que reportó el cliente y que motivan la "
        "reingeniería:",
    )

    add_heading(doc, "4.1 Operativos", level=2)
    add_bullets(
        doc,
        [
            "Sin sincronización entre sucursales: cambios de precio se replican manualmente, lo que genera diferencias y errores en el mostrador.",
            "Compras a proveedor totalmente manuales: cargar OC, recibir mercadería, conciliar contra factura del proveedor — todo en papel o en Excel.",
            "OCR de comprobantes inexistente: cuando llega una factura del proveedor en papel, alguien la tipea a mano en el sistema.",
            "Sin acceso remoto: para consultar algo fuera del local hay que conectarse al servidor por VPN (que no existe) o ir físicamente.",
        ],
    )

    add_heading(doc, "4.2 Estratégicos", level=2)
    add_bullets(
        doc,
        [
            "Ningún reporting moderno: para entender qué productos rotan más, en qué horario, en qué sucursal, hay que armar planillas manualmente.",
            "Sin alertas proactivas: stock bajo, productos por vencer, faltantes — todo se descubre tarde.",
            "Calendario financiero a ojo: cheques, tarjetas y compromisos de pago se llevan en cuaderno.",
        ],
    )

    add_heading(doc, "4.3 Killer feature solicitada", level=2)
    add_paragraph(
        doc,
        "El requerimiento más prioritario del cliente — la \"killer feature\" "
        "que justifica la inversión — es la conciliación tri-lateral de "
        "compras: para cada compra, comparar la Orden de Compra emitida, la "
        "Recepción física de mercadería (con scanner de código de barras) y "
        "la Factura del proveedor (cargada por OCR), y mostrar diferencias "
        "con un semáforo verde / amarillo / rojo. Esto reemplaza el trabajo "
        "del empleado administrativo dedicado a controlar compras.",
        bold=True,
    )

    doc.add_page_break()

    # ================== 5. Visión y objetivos ==================
    add_heading(doc, "5. Visión y objetivos", level=1)

    add_heading(doc, "5.1 Visión", level=2)
    add_paragraph(
        doc,
        "Plataforma única de gestión integral para CASA SALCO que centralice "
        "datos de las 4 sucursales en tiempo real, automatice tareas "
        "administrativas (compras, conciliación, OCR, AFIP), y ofrezca "
        "información operativa y estratégica accesible desde cualquier "
        "lugar y dispositivo.",
    )

    add_heading(doc, "5.2 Objetivos cuantificables", level=2)
    add_bullets(
        doc,
        [
            "Reducir el tiempo de carga manual de facturas de compra mediante OCR (objetivo: < 30 segundos por comprobante).",
            "Eliminar diferencias de precios entre sucursales (sync online en menos de 1 segundo).",
            "Sustituir un puesto administrativo dedicado a control de compras con la conciliación tri-lateral automática.",
            "Permitir consulta y operación desde dispositivos móviles para repositores y dueños.",
            "Cumplir con régimen fiscal argentino: facturación electrónica AFIP (WSFEv1), libro IVA digital, retenciones.",
        ],
    )

    add_heading(doc, "5.3 Fuera de alcance del PMV", level=2)
    add_bullets(
        doc,
        [
            "Tienda online / e-commerce.",
            "App de fidelización para clientes finales.",
            "Integración con sistema contable externo (Tango, Bejerman, etc.). Se exporta a Excel para que el contador procese.",
            "Manejo de sucursales en distintas razones sociales — todas las sucursales operan bajo CASA SALCO.",
        ],
    )

    doc.add_page_break()

    # ================== 6. PMV — Requerimientos funcionales ==================
    add_heading(doc, "6. Requerimientos funcionales del PMV", level=1)

    add_paragraph(
        doc,
        "30 requerimientos funcionales (RF-01..RF-30) organizados en 9 "
        "módulos. Surgen del relevamiento del cliente y de la observación "
        "del sistema viejo. Cada RF se mapea contra la implementación en "
        "el documento Gap-Producción.",
    )

    add_heading(doc, "6.1 Productos (RF-01..RF-05)", level=2)
    add_rf_table(
        doc,
        [
            ("RF-01", "Productos", "Maestro de artículos con múltiples códigos de barras (principal + alternativos por proveedor + empaquetados)."),
            ("RF-02", "Productos", "Soporte de unidad base + subunidades / equivalencias (caja x N unidades, etc.)."),
            ("RF-03", "Productos", "Productos fraccionables (fiambres, panadería, productos pesables) con paso de fracción."),
            ("RF-04", "Productos", "Importación / exportación masiva desde Excel."),
            ("RF-05", "Productos", "Historial de cambios de precios y de proveedor visualizable como eventos."),
        ],
    )

    add_heading(doc, "6.2 Proveedores (RF-06..RF-07)", level=2)
    add_rf_table(
        doc,
        [
            ("RF-06", "Proveedores", "Datos de contacto + frecuencia de visita (semanal / quincenal / mensual)."),
            ("RF-07", "Proveedores", "Catálogo de productos por proveedor con presentaciones (caja, unidad, x12, etc.) y costos históricos."),
        ],
    )

    add_heading(doc, "6.3 Compras — KILLER FEATURE (RF-08..RF-11)", level=2)
    add_rf_table(
        doc,
        [
            ("RF-08", "Compras", "Orden de Compra con sugerencia de cantidades por reposición."),
            ("RF-09", "Compras", "Recepción de mercadería con escáner de código de barras."),
            ("RF-10", "Compras", "CONCILIACIÓN TRI-LATERAL: OC ↔ Recepción ↔ Factura (OCR), con semáforo verde/amarillo/rojo."),
            ("RF-11", "Compras", "Matching inteligente de ítems por descripción / código del proveedor / historial de compras."),
        ],
    )

    add_heading(doc, "6.4 OCR de comprobantes (RF-12..RF-14)", level=2)
    add_rf_table(
        doc,
        [
            ("RF-12", "OCR", "Carga de imagen / PDF de comprobante de compra desde la UI."),
            ("RF-13", "OCR", "Extracción automática de proveedor, fecha, productos, cantidades y precios mediante IA (Claude Vision o Gemini)."),
            ("RF-14", "OCR", "Validación asistida: el operador revisa los datos extraídos lado-a-lado con la imagen original antes de confirmar."),
        ],
    )

    add_heading(doc, "6.5 Stock (RF-15..RF-17)", level=2)
    add_rf_table(
        doc,
        [
            ("RF-15", "Stock", "Stock en tiempo real por sucursal."),
            ("RF-16", "Stock", "Movimientos de entrada / salida / ajuste con motivo."),
            ("RF-17", "Stock", "Alertas proactivas: bajo stock + vencimientos próximos."),
        ],
    )

    add_heading(doc, "6.6 POS / Ventas (RF-18..RF-22)", level=2)
    add_rf_table(
        doc,
        [
            ("RF-18", "POS", "Búsqueda instantánea tipo Odoo (debounce + filtros inline)."),
            ("RF-19", "POS", "Integración con balanza Systel / Kretz para artículos pesables."),
            ("RF-20", "POS", "Pago en efectivo, tarjeta crédito/débito, QR (Mercado Pago / MODO), cheques, cuenta corriente, vales."),
            ("RF-21", "POS", "Integración con Mercado Pago / Clover / Payway que evite la entrada manual del monto (webhook + match)."),
            ("RF-22", "POS", "Cierre de caja con totales por medio de pago y reporte de diferencias."),
        ],
    )

    add_heading(doc, "6.7 Dashboard / Reporting (RF-23..RF-25)", level=2)
    add_rf_table(
        doc,
        [
            ("RF-23", "Dashboard", "Home con KPIs: ventas del día, caja, alertas, vencimientos."),
            ("RF-24", "Dashboard", "Calendario financiero: pagos a proveedores + tarjetas + acreditaciones de tarjetas + flujo proyectado."),
            ("RF-25", "Dashboard", "Paneles diferenciados por rol (dueño / cajero / repositor / contador)."),
        ],
    )

    add_heading(doc, "6.8 Acceso remoto y mobile (RF-26..RF-27, RF-30)", level=2)
    add_rf_table(
        doc,
        [
            ("RF-26", "Remoto", "Web app accesible desde cualquier navegador, multi-sucursal, con autenticación por roles."),
            ("RF-27", "Remoto", "Mobile-friendly (responsive + PWA instalable)."),
            ("RF-30", "Mobile", "App ligera para repositor: escaneo + recepción + consulta de stock + alertas."),
        ],
    )

    add_heading(doc, "6.9 Usuarios e importación (RF-28..RF-29)", level=2)
    add_rf_table(
        doc,
        [
            ("RF-28", "Usuarios", "Roles jerárquicos: admin, supervisor, cajero, fiambrero, repositor, contador (rol \"compras\" mapeado a supervisor)."),
            ("RF-29", "Importación", "Importador masivo de productos + precios + proveedores desde el formato legacy del cliente (xls BIFF cp1252)."),
        ],
    )

    doc.add_page_break()

    # ================== 7. Restricciones ==================
    add_heading(doc, "7. Restricciones", level=1)

    add_heading(doc, "7.1 Normativas (Argentina)", level=2)
    add_bullets(
        doc,
        [
            "Facturación electrónica AFIP — cumplimiento del régimen WSFEv1 vía PyAfipWs.",
            "Libro IVA digital — exportable para presentación.",
            "Régimen de retenciones nacionales y provinciales según corresponda.",
            "Conservación de comprobantes (físicos digitalizados via OCR + originales en papel donde aplique).",
        ],
    )

    add_heading(doc, "7.2 Hardware existente del cliente", level=2)
    add_bullets(
        doc,
        [
            "Impresoras térmicas 3NSTAR ESC/POS (tickets de venta).",
            "Balanzas Systel y Kretz con interfaz serial (artículos pesables).",
            "PCs de mostrador (Windows) — ya instaladas y operativas.",
            "Conexión a internet en cada sucursal (con caídas ocasionales — el sistema debe tolerar offline).",
        ],
    )

    add_heading(doc, "7.3 Operativas / de equipo", level=2)
    add_bullets(
        doc,
        [
            "Equipo de desarrollo: 1 desarrollador. Implica diseñar para mantenibilidad por una persona — descarta arquitecturas tipo microservicios.",
            "Cliente prefiere caminos incrementales y seguros frente a entregas grandes (\"menor costo y más seguro\").",
            "Operaciones destructivas (DROP, DELETE masivo, migraciones que pisan datos) requieren confirmación explícita del cliente con backup previo.",
            "Demos al cliente periódicas — el cliente acompaña el desarrollo y ve avances funcionales.",
        ],
    )

    doc.add_page_break()

    # ================== 8. Decisiones arquitectónicas ==================
    add_heading(doc, "8. Decisiones arquitectónicas iniciales", level=1)

    add_paragraph(
        doc,
        "Decisiones tomadas al inicio del proyecto, documentadas como ADRs "
        "(Architecture Decision Records) en `docs/architecture.md` del repo.",
    )

    add_heading(doc, "8.1 Stack", level=2)
    add_bullets(
        doc,
        [
            "Backend: Flask (Python) + SQLAlchemy + Alembic + Flask-SocketIO + Celery.",
            "Validación: Pydantic v2.",
            "Base de datos: PostgreSQL central + SQLite local en cada caja (offline-tolerant).",
            "Cache / pub-sub: Redis.",
            "Frontend: React + Vite + TypeScript + Tailwind + shadcn/ui.",
            "Estética: Apple-like (SF fonts, Apple Blue accent, radius 12px, glass surfaces).",
            "POS desktop: Tauri 2 + React (mismo bundle que la web).",
            "Mobile: PWA con la misma codebase.",
            "AFIP: PyAfipWs (biblioteca argentina madura).",
            "Hardware: agente Python sidecar (Flask local) que habla con balanzas y impresoras vía serial / USB.",
            "Infra de deploy: Railway o Fly.io para el backend central.",
        ],
    )

    add_heading(doc, "8.2 Principios", level=2)
    add_bullets(
        doc,
        [
            "Backend monolito modular (blueprints) — no microservicios, dado que se mantiene con 1 dev.",
            "Local-first para cajas: si internet se cae, la venta sigue. Cola local que se drena al reconectar. AFIP es la única operación que bloquea sin internet (no hay CAEA en MVP).",
            "Una sola UI React: misma codebase se sirve por Flask (web), corre en Tauri (caja desktop) y se instala como PWA (mobile).",
            "Sync en vivo: Flask-SocketIO para precios; cambios se propagan a todas las sucursales en menos de 1s.",
            "Observabilidad: logs estructurados, errores a Sentry (en producción).",
            "Seguridad: JWT con refresh tokens, roles jerárquicos, validación back+front, defense-in-depth.",
        ],
    )

    add_heading(doc, "8.3 Topología objetivo (Fase 2)", level=2)
    add_paragraph(
        doc,
        "Servidor central en cloud (Railway/Fly) con Postgres + Redis. Cada "
        "sucursal corre su POS Tauri con SQLite local. La sincronización es "
        "vía HTTPS + WebSocket: cambios del central llegan en vivo, "
        "operaciones locales se encolan y suben cuando hay conexión. El "
        "dashboard web y la PWA mobile consumen la misma API.",
    )

    doc.add_page_break()

    # ================== 9. Roadmap por fases ==================
    add_heading(doc, "9. Roadmap por fases", level=1)

    add_paragraph(
        doc,
        "El alcance se divide en 4 fases progresivas para entregar valor "
        "temprano y validar con el cliente antes de avanzar.",
    )

    add_heading(doc, "Fase 1 — MVP (8-10 semanas)", level=2)
    add_paragraph(doc, "Objetivo: catálogo unificado y sincronización de precios entre sucursales en tiempo real.")
    add_bullets(
        doc,
        [
            "ETL desde DBFs y .xls legacy (RF-04, RF-29).",
            "CRUD de artículos, clientes, proveedores (RF-01..RF-07).",
            "Sincronización de precios con WebSocket (Dolor #1 del cliente).",
            "Dashboard básico con KPIs (RF-23 parcial).",
        ],
    )

    add_heading(doc, "Fase 2 — POS y AFIP (10-12 semanas)", level=2)
    add_paragraph(doc, "Objetivo: operar mostrador y emitir comprobantes fiscales.")
    add_bullets(
        doc,
        [
            "POS Tauri con balanza, impresora, búsqueda instantánea (RF-18, RF-19).",
            "Integración AFIP con WSFEv1 (cert + key + emisión de CAE) (RF-22).",
            "Múltiples medios de pago — efectivo, tarjeta, QR (RF-20).",
            "Cierre de caja con arqueo (RF-22).",
            "Cola offline: ventas que se sincronizan al reconectar.",
        ],
    )

    add_heading(doc, "Fase 3 — Dashboard premium y PWA (8 semanas)", level=2)
    add_paragraph(doc, "Objetivo: información estratégica y acceso remoto.")
    add_bullets(
        doc,
        [
            "Dashboard avanzado: heatmaps, mapas, correlaciones, top productos (RF-23, RF-24).",
            "Alertas proactivas: stock bajo + vencimientos (RF-17).",
            "Calendario financiero — compromisos de pago, tarjetas (RF-24).",
            "PWA mobile responsive (RF-27).",
            "Layout diferenciado por rol (RF-25).",
        ],
    )

    add_heading(doc, "Fase 4 — OCR y datamining (6-8 semanas)", level=2)
    add_paragraph(doc, "Objetivo: automatizar la carga de comprobantes de compra y mostrar inteligencia de datos.")
    add_bullets(
        doc,
        [
            "OCR de comprobantes con Claude Vision / Gemini (RF-12..RF-14).",
            "Conciliación tri-lateral OC ↔ Recepción ↔ Factura — KILLER (RF-08..RF-11).",
            "Mobile app de repositor: escanear, recibir, consultar (RF-30).",
            "Export Excel para contador (libro IVA, retenciones, cuentas corrientes).",
            "Datamining: análisis de canasta, productos correlacionados, sugerencias.",
        ],
    )

    doc.add_page_break()

    # ================== 10. Riesgos identificados ==================
    add_heading(doc, "10. Riesgos identificados al inicio", level=1)

    add_paragraph(doc, "Riesgos detectados durante el relevamiento, con su mitigación inicial:")

    add_bullets(
        doc,
        [
            "Datos legacy mal codificados — los .xls vienen en BIFF/cp1252 con strings mal codificados. Mitigación: re-decode en Python al importar y reporte de saltadas.",
            "Datos sucios en el legacy — filas de prueba, productos discontinuos, registros incompletos. Mitigación: validación + filtro al importar con reporte detallado para el cliente.",
            "Hardware fiscal AFIP — el cert dura 2 años; renovación crítica. Mitigación: agendar renewal y documentar el procedimiento.",
            "Integraciones de pago (Mercado Pago / Clover / Payway) tienen sandboxes separados. Mitigación: Fase 4, comenzar por una sola y agregar gradualmente.",
            "Conexión a internet inestable en sucursales — cortes ocasionales rompen las operaciones. Mitigación: arquitectura local-first con cola offline.",
            "Volumen del catálogo — 35.000 artículos en una SQLite local pueden generar lock contention. Mitigación: Postgres central en producción + índices apropiados.",
            "Hardware POS específico — balanzas Systel / Kretz con protocolos serial propietarios. Mitigación: agente Python con drivers + modo mock para development.",
            "Productos vencidos / lotes — sin manejo de lote en el viejo, los vencimientos se descubren tarde. Mitigación: tabla LoteStock con fecha_vencimiento, alertas proactivas (Fase 3).",
            "Mantenibilidad por 1 dev — riesgo de bus factor. Mitigación: documentación, tests, memoria persistente de decisiones.",
            "Seguridad — gestión de roles, secrets, datos fiscales. Mitigación: defense-in-depth (validación back+front), JWT con expiración, secrets reales en producción, auditoría futura.",
        ],
    )

    doc.add_page_break()

    # ================== 11. Glosario ==================
    add_heading(doc, "11. Glosario", level=1)

    add_bullets(
        doc,
        [
            "AFIP — Administración Federal de Ingresos Públicos (organismo recaudador argentino).",
            "WSFEv1 — Web Service de Facturación Electrónica de AFIP, versión 1.",
            "CAE — Código de Autorización Electrónica que AFIP devuelve al autorizar una factura.",
            "PyAfipWs — biblioteca Python para integración con AFIP.",
            "BIFF — Binary Interchange File Format, formato de Excel pre-2007 (.xls).",
            "DBF — formato de base de datos de Clipper / dBase / FoxPro / Harbour.",
            "OC — Orden de Compra.",
            "Conciliación tri-lateral — comparar OC, recepción y factura del proveedor para detectar diferencias.",
            "PMV — Producto Mínimo Viable.",
            "PWA — Progressive Web App (instalable, offline-tolerant).",
            "ESC/POS — protocolo de comandos para impresoras térmicas.",
            "MODO — billetera virtual de bancos argentinos.",
            "Dragonfish — software ERP comercial argentino con UX dado por anti-pattern del proyecto.",
        ],
    )

    # ================== Cierre ==================
    doc.add_paragraph()
    add_paragraph(
        doc,
        f"Documento generado el {today}. Este relevamiento es la base de la "
        "que nacieron los 30 requerimientos funcionales del PMV (RF-01..RF-30) "
        "y las decisiones arquitectónicas iniciales. Para ver el estado de "
        "cumplimiento de cada RF y los pendientes a producción, consultar el "
        "documento `GAP-Producción.docx`.",
        italic=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK: {OUT.relative_to(ROOT)} ({OUT.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
