"""Genera 4 manuales .docx (uno por rol) con texto + screenshots embebidos.

Pre-requisito: correr `capture_screenshots.py` antes para tener las imágenes.

Uso:
    cd backend
    .venv/Scripts/python ../../scripts/generate_manuals.py

Output: `manuals/output/manual-<rol>.docx`
"""
from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SCREENSHOTS_DIR = ROOT / "manuals" / "screenshots"
OUTPUT_DIR = ROOT / "manuals" / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

ROLE_DISPLAY = {
    "admin": "Administrador",
    "supervisor": "Supervisor",
    "cajero": "Cajero",
    "contador": "Contador",
}

ROLE_CREDS = {
    "admin": ("admin@casasalco.app", "admin123"),
    "supervisor": ("supervisor1@casasalco.app", "super123"),
    "cajero": ("cajero1@casasalco.app", "cajero123"),
    "contador": ("contador@casasalco.app", "contador123"),
}

# Resumen del rol — qué hace, principales responsabilidades.
ROLE_DESCRIPTION = {
    "admin": (
        "El rol Administrador es el de mayor jerarquía del sistema. Tiene acceso "
        "completo a todos los módulos: operación diaria, backoffice, configuración "
        "y mantenimiento de catálogos. Es el único que puede crear/eliminar "
        "usuarios, modificar parámetros de sucursal y administrar el sistema."
    ),
    "supervisor": (
        "El Supervisor administra la operación diaria del comercio: controla "
        "ventas, gestiona artículos y proveedores, supervisa stock y autoriza "
        "movimientos. No puede modificar configuración de sucursales ni "
        "administrar usuarios — eso queda reservado al Administrador."
    ),
    "cajero": (
        "El Cajero opera el día a día en mostrador: registra ventas en el Punto "
        "de venta, consulta facturas y movimientos, mira stock y atiende "
        "clientes. No accede a información sensible de proveedores, compras, "
        "ni reportes contables."
    ),
    "contador": (
        "El Contador trabaja con la información financiera y fiscal: factura, "
        "movimientos, cuentas a pagar, proveedores, compras, consultas y "
        "reportes para AFIP. No opera el Punto de venta ni administra stock."
    ),
}

# Pantallas — descripción y acciones por cada path.
SCREEN_META: dict[str, dict] = {
    "/": {
        "slug": "01-dashboard",
        "title": "Dashboard",
        "desc": (
            "Es la primera pantalla que ves al iniciar sesión. Muestra un panorama "
            "general del estado del comercio: ventas del día, alertas activas, "
            "stock crítico y accesos rápidos a las funciones más usadas."
        ),
        "acciones": [
            "Revisar KPIs del día (ventas, ticket promedio, etc.)",
            "Ver alertas activas (stock bajo, vencimientos, etc.)",
            "Acceder rápidamente al Punto de venta",
            "Consultar últimas facturas emitidas",
        ],
    },
    "/pos": {
        "slug": "02-pos",
        "title": "Punto de venta (POS)",
        "desc": (
            "Pantalla principal de operación de mostrador. Permite cargar artículos "
            "(por código de barras, búsqueda manual o teclado), aplicar descuentos, "
            "seleccionar cliente y emitir el comprobante. Optimizada para ser "
            "rápida y precisa con teclado."
        ),
        "acciones": [
            "Buscar y agregar artículos al carrito",
            "Modificar cantidades y aplicar descuentos por línea",
            "Seleccionar cliente y forma de pago",
            "Emitir factura A / B / C / X según corresponda",
            "Anular o pausar venta",
        ],
        "tips": [
            "Atajos de teclado para operar sin mouse.",
            "Funciona offline: las ventas se encolan y se sincronizan cuando vuelve la conexión.",
        ],
    },
    "/facturas": {
        "slug": "03-facturas",
        "title": "Facturas",
        "desc": (
            "Histórico de comprobantes emitidos. Permite filtrar por fecha, "
            "cliente, tipo (A/B/C/X), estado AFIP y descargar/reimprimir."
        ),
        "acciones": [
            "Filtrar por fecha, cliente, tipo o estado",
            "Ver detalle de un comprobante",
            "Reimprimir factura / ticket",
            "Anular comprobante (con motivos y permiso correspondiente)",
        ],
    },
    "/movimientos": {
        "slug": "04-movimientos",
        "title": "Movimientos de caja",
        "desc": (
            "Registro de todos los ingresos y egresos del día por caja. Incluye "
            "ventas, cobranzas, gastos y movimientos manuales. Soporta arqueo."
        ),
        "acciones": [
            "Ver movimientos del día por caja",
            "Cargar ingresos/egresos manuales",
            "Realizar arqueo / cierre de caja",
            "Filtrar por tipo de movimiento (efectivo, tarjeta, transferencia)",
        ],
    },
    "/pagos": {
        "slug": "05-pagos",
        "title": "Calendario de pagos",
        "desc": (
            "Vista unificada de compromisos de pago: facturas de proveedores, "
            "tarjetas, servicios. Permite planificar el cash-flow."
        ),
        "acciones": [
            "Ver compromisos próximos en formato calendario o lista",
            "Marcar pago como realizado",
            "Generar compromisos automáticamente desde facturas de proveedor",
            "Administrar tarjetas y vencimientos",
        ],
    },
    "/articulos": {
        "slug": "06-articulos",
        "title": "Artículos",
        "desc": (
            "Catálogo maestro de productos. Búsqueda por código, descripción, "
            "familia, marca o proveedor. Soporta múltiples códigos por artículo "
            "(EAN principal, alternos, empaquetados)."
        ),
        "acciones": [
            "Buscar artículos por código o descripción",
            "Filtrar por familia / rubro / marca / proveedor",
            "Ver detalle: precios, stock, presentaciones",
            "Crear / editar / eliminar artículos (según rol)",
        ],
    },
    "/clientes": {
        "slug": "07-clientes",
        "title": "Clientes",
        "desc": (
            "Padrón de clientes registrados. Permite alta de cliente con CUIT, "
            "ver cuenta corriente, historial de facturas y datos de contacto."
        ),
        "acciones": [
            "Buscar / filtrar clientes",
            "Crear nuevo cliente",
            "Ver detalle, cuenta corriente, historial",
            "Editar datos del cliente",
        ],
    },
    "/proveedores": {
        "slug": "08-proveedores",
        "title": "Proveedores",
        "desc": (
            "Padrón de proveedores. Cada proveedor puede tener múltiples "
            "artículos asociados con códigos y presentaciones específicas."
        ),
        "acciones": [
            "Buscar / filtrar proveedores",
            "Crear nuevo proveedor",
            "Ver artículos asociados al proveedor",
            "Editar datos y cuenta corriente",
        ],
    },
    "/compras/ocr": {
        "slug": "09-compras",
        "title": "Compras (OCR de comprobantes)",
        "desc": (
            "Carga de facturas de compra con OCR — sacás foto del comprobante "
            "del proveedor y la IA extrae los datos automáticamente: proveedor, "
            "fecha, ítems, totales. Después confirmás y se carga en el sistema."
        ),
        "acciones": [
            "Subir foto/PDF de factura de compra",
            "Revisar y editar los datos extraídos",
            "Confirmar la creación de la factura",
            "Ver historial de comprobantes procesados",
        ],
    },
    "/stock": {
        "slug": "10-stock",
        "title": "Stock",
        "desc": (
            "Listado de stock por sucursal. Indicadores por estado (agotado, "
            "crítico, en reorden, OK, sobrestock). Búsqueda por código/descripción "
            "y ajuste manual con motivo."
        ),
        "acciones": [
            "Filtrar por estado de reposición",
            "Buscar artículos específicos",
            "Ajustar cantidad manual con motivo",
            "Configurar mínimos / máximos / punto de reorden por artículo",
        ],
        "tips": [
            "Los valores 'efectivos' usan el override de sucursal o, si no, el default del artículo.",
        ],
    },
    "/stock/reposicion": {
        "slug": "11-reposicion",
        "title": "Reposición",
        "desc": (
            "Sugerencias de compra agrupadas por proveedor. El sistema calcula "
            "qué pedir según punto de reorden, stock máximo y velocidad de venta."
        ),
        "acciones": [
            "Ver sugerencias agrupadas por proveedor",
            "Recalcular stock óptimo basado en histórico de ventas",
            "Generar orden de compra desde la sugerencia",
            "Exportar lista de reposición a Excel",
        ],
    },
    "/sucursales": {
        "slug": "12-sucursales",
        "title": "Sucursales",
        "desc": (
            "Administración de sucursales del comercio. Cada sucursal tiene sus "
            "áreas (Comestibles, Drugstore, etc.), datos de contacto y geolocalización."
        ),
        "acciones": [
            "Ver listado de sucursales activas",
            "Crear / editar / desactivar sucursales",
            "Configurar áreas internas de la sucursal",
            "Definir datos fiscales y de contacto",
        ],
    },
    "/consultas": {
        "slug": "13-consultas",
        "title": "Consultas (BI)",
        "desc": (
            "Equivalente al F3 del sistema viejo. Consultas analíticas sobre "
            "ventas, stock, cobranzas, pagos. Permite filtrar y exportar."
        ),
        "acciones": [
            "Elegir entidad (ventas, stock, clientes, etc.)",
            "Aplicar filtros (fecha, sucursal, etc.)",
            "Ver tabla con resultados",
            "Exportar a Excel para análisis offline",
        ],
    },
    "/exports": {
        "slug": "14-exports",
        "title": "Exportar",
        "desc": (
            "Centro de exportación de reportes en Excel: ventas, stock valorizado, "
            "cuentas corrientes, libro IVA digital, resúmenes para AFIP."
        ),
        "acciones": [
            "Reportes fiscales (Libro IVA digital)",
            "Reportes comerciales (ventas, cobranzas)",
            "Reportes de stock (listado, valorizado)",
            "Cuentas corrientes (clientes y proveedores)",
        ],
    },
    "/mantenimiento": {
        "slug": "15-mantenimiento",
        "title": "Mantenimiento de catálogos",
        "desc": (
            "Administración de catálogos maestros: Familias, Rubros, Subrubros y "
            "Marcas. Estos catálogos clasifican los artículos jerárquicamente."
        ),
        "acciones": [
            "Crear / editar / eliminar familias y rubros",
            "Administrar subrubros (anidados en rubros)",
            "Gestionar marcas",
            "Mantener consistencia del catálogo",
        ],
    },
}

ROLE_SCREENS: dict[str, list[str]] = {
    "admin": list(SCREEN_META.keys()),
    "supervisor": [
        p for p in SCREEN_META if p not in ("/sucursales", "/mantenimiento")
    ],
    "cajero": ["/", "/pos", "/facturas", "/movimientos", "/articulos", "/clientes", "/stock"],
    "contador": [
        "/",
        "/facturas",
        "/movimientos",
        "/pagos",
        "/articulos",
        "/clientes",
        "/proveedores",
        "/compras/ocr",
        "/consultas",
        "/exports",
    ],
}


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    if level == 0:
        for run in h.runs:
            run.font.size = Pt(28)


def add_paragraph(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_screenshot(doc: Document, path: Path, *, width_cm: float = 16.0) -> None:
    if not path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Screenshot no disponible: {path.name}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        return
    try:
        doc.add_picture(str(path), width=Cm(width_cm))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f"[Error al insertar imagen {path.name}: {e}]")
        run.italic = True


def build_manual(role: str) -> Path:
    doc = Document()
    role_human = ROLE_DISPLAY[role]
    email, password = ROLE_CREDS[role]
    today = dt.date.today().isoformat()

    # ---- Portada ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Manual de Usuario")
    run.font.size = Pt(34)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Rol: {role_human}")
    run.font.size = Pt(20)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("CASA SALCO ERP")
    run.font.size = Pt(14)
    run.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generado: {today}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ---- Introducción ----
    add_heading(doc, "Introducción", level=1)
    add_paragraph(doc, ROLE_DESCRIPTION[role])
    doc.add_paragraph()

    add_heading(doc, "Cómo iniciar sesión", level=2)
    add_paragraph(
        doc,
        "Abrí tu navegador (Chrome, Edge o Firefox) y entrá a la URL del sistema. "
        "Ingresá tus credenciales y presioná Iniciar sesión.",
    )
    p = doc.add_paragraph()
    p.add_run("URL: ").bold = True
    p.add_run("https://laagencia.myvnc.com/casasalco/")
    p = doc.add_paragraph()
    p.add_run("Usuario: ").bold = True
    p.add_run(email)
    p = doc.add_paragraph()
    p.add_run("Contraseña: ").bold = True
    p.add_run(password)

    doc.add_paragraph()
    add_paragraph(
        doc,
        "Si olvidás tu contraseña, contactá al Administrador del sistema. "
        "Si entrás desde un dispositivo nuevo, podés instalar la app como PWA "
        "para acceder más rápido (botón 'Instalar' en el navegador).",
        italic=True,
    )

    doc.add_page_break()

    # ---- Una sección por pantalla ----
    add_heading(doc, "Pantallas y funciones", level=1)
    add_paragraph(
        doc,
        "A continuación se documenta cada pantalla a la que tenés acceso, qué "
        "hace y qué acciones podés realizar. Los menús que no aparecen en tu "
        "barra lateral están disponibles para otros roles del sistema.",
    )

    for path in ROLE_SCREENS[role]:
        meta = SCREEN_META[path]
        slug = meta["slug"]
        screenshot = SCREENSHOTS_DIR / role / f"{slug}.png"

        doc.add_paragraph()
        add_heading(doc, meta["title"], level=2)
        add_paragraph(doc, meta["desc"])

        # Screenshot
        if screenshot.exists():
            add_screenshot(doc, screenshot)
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(f"Pantalla: {meta['title']}")
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Acciones
        if meta.get("acciones"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Qué podés hacer:").bold = True
            add_bullets(doc, meta["acciones"])

        # Tips
        if meta.get("tips"):
            p = doc.add_paragraph()
            p.add_run("Tips:").bold = True
            add_bullets(doc, meta["tips"])

    # ---- Cierre ----
    doc.add_page_break()
    add_heading(doc, "Soporte y ayuda", level=1)
    add_paragraph(
        doc,
        "Si tenés un problema operativo o detectás un error, comunicate con el "
        "Administrador del sistema. Antes de reportar, anotá: qué pantalla "
        "estabas usando, qué acción intentaste y qué mensaje de error apareció. "
        "Eso acelera la resolución.",
    )
    doc.add_paragraph()
    add_paragraph(
        doc,
        f"Manual generado automáticamente el {today} desde la versión productiva "
        "del sistema. Si encontrás diferencias entre lo que muestra el manual y "
        "la app real, prevalece la app — pueden haber actualizaciones recientes.",
        italic=True,
    )

    out_path = OUTPUT_DIR / f"manual-{role}.docx"
    doc.save(str(out_path))
    return out_path


def main() -> int:
    for role in ROLE_DISPLAY:
        out = build_manual(role)
        size_kb = out.stat().st_size // 1024
        print(f"[{role}] {out.name} ({size_kb} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
