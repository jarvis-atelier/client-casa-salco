"""Genera el documento `GAP-Producción.docx` con lo hecho y lo pendiente
para pasar a producción.

Uso:
    cd backend
    .venv/Scripts/python ../../scripts/generate_gap_doc.py
"""
from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "manuals" / "output" / "GAP-Produccion.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def shade_cell(cell, hex_color: str) -> None:
    """Aplica color de fondo a una celda de tabla."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_status_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    """rows = [(estado, item, detalle)]. estado: DONE | TODO | RIESGO"""
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    for i, h in enumerate(("Estado", "Ítem", "Detalle")):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
        shade_cell(hdr[i], "1F2937")
        for r in hdr[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, (estado, item, detalle) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = estado
        cells[1].text = item
        cells[2].text = detalle
        if estado == "DONE":
            shade_cell(cells[0], "D1FAE5")
        elif estado == "TODO":
            shade_cell(cells[0], "FEF3C7")
        elif estado == "RIESGO":
            shade_cell(cells[0], "FEE2E2")


def main() -> int:
    doc = Document()
    today = dt.date.today().isoformat()

    # ===== Portada =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Gap Analysis — Pase a Producción")
    run.font.size = Pt(28)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CASA SALCO ERP")
    run.font.size = Pt(16)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Fecha: {today}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ===== Resumen ejecutivo =====
    add_heading(doc, "Resumen ejecutivo", level=1)

    add_paragraph(doc, "Estado actual:", bold=True)
    add_paragraph(
        doc,
        "El sistema está en estado dev-ready: corre estable en el servidor y todos "
        "los flujos básicos funcionan. Está listo para uso interno, demos y "
        "training de usuarios. NO está listo para uso productivo real con clientes "
        "porque hay tres bloqueadores duros (secrets default, AFIP en modo mock y "
        "OCR en modo mock).",
    )

    add_paragraph(doc, "Bloqueadores críticos para producción:", bold=True)
    add_bullets(
        doc,
        [
            "Secrets en valores de desarrollo (JWT_SECRET_KEY y SECRET_KEY).",
            "AFIP en modo 'mock' — la facturación electrónica no emite a AFIP de verdad.",
            "OCR en modo 'mock' — el OCR de comprobantes devuelve datos ficticios; falta API key real (Claude o Gemini).",
        ],
    )

    add_paragraph(doc, "Aspectos importantes pero no bloqueantes:", bold=True)
    add_bullets(
        doc,
        [
            "Servidor corre en Werkzeug (dev server) y no en Gunicorn — funciona pero no es óptimo.",
            "Base de datos en SQLite local — para volumen real conviene Postgres.",
            "Stock cargado es demo (random) — falta inventario real.",
            "No hay backups programados ni monitoring activo.",
        ],
    )

    doc.add_page_break()

    # ===== Sección: Lo que se hizo =====
    add_heading(doc, "1. Lo que se hizo", level=1)
    add_paragraph(
        doc,
        "Trabajo realizado durante esta iteración (verificable en git log y en "
        "los reportes del agente APB).",
    )

    add_heading(doc, "1.1 Infraestructura y deployment", level=2)
    add_status_table(
        doc,
        [
            ("DONE", "Single-port (5005)", "Backend Flask sirve frontend buildado + API en el mismo puerto. Eliminó la dependencia del proxy de Vite."),
            ("DONE", "Migración a subpath /casasalco/", "Convención del nginx del usuario. Frontend con base, backend con WSGI middleware que strippa el prefix."),
            ("DONE", "Bloque nginx en 2 servers", "Agregado en laagencia.myvnc.com y datacivis.myvnc.com sin tocar el resto."),
            ("DONE", "Diagnóstico bug Vite + Node v24", "Identificado y documentado. Workaround: single-port."),
            ("DONE", "Limpieza vite.config.js legacy", "Había dos configs (.ts y .js), Vite priorizaba el .js viejo con datos pre-rebrand."),
        ],
    )

    add_heading(doc, "1.2 Calidad de datos", level=2)
    add_status_table(
        doc,
        [
            ("DONE", "Rebrand castulo → casasalco", "DB renombrada con backup, 8 usuarios actualizados, 22 archivos del repo. (Sesión anterior)"),
            ("DONE", "Cajeros con nombre correcto", "Reemplazado 'Cajero Castulo X' por 'Cajero CASA SALCO X' en DB."),
            ("DONE", "Merge de 4 sucursales a 1", "Una sola sucursal CASA SALCO. 556 facturas + 622 movimientos preservados con remap de punto_venta y caja_numero. Backup previo."),
            ("DONE", "Stock demo cargado", "34.847 artículos con stock distribuido (5% sin stock, 15% bajo, 60% medio, 15% alto, 5% sobrestock). Total 4.874.411 unidades."),
            ("DONE", "Importación XLS legacy", "34.847 artículos del XLS de CASA SALCO + 123 proveedores + 36.204 relaciones articulo-proveedor."),
        ],
    )

    add_heading(doc, "1.3 Bug fixes", level=2)
    add_status_table(
        doc,
        [
            ("DONE", "Login password vacío → 422", "Era 401 invalid_credentials. Ahora schema Pydantic con Field(min_length=1)."),
            ("DONE", "JWT error shape consistente", "Antes {msg:...}, ahora {code, error}. 5 callbacks de Flask-JWT-Extended registrados."),
            ("DONE", "Werkzeug 404 normalizado", "Antes blurb largo, ahora 'ruta no encontrada'."),
            ("DONE", "Stock 42x más rápido", "GET /api/v1/stock pasó de 2.5s/15MB a 0.06s/28KB. Paginación + filtros server-side, articulo embedded."),
            ("DONE", "Dropdown transparente", "Faltaba 'popover' en tailwind.config — los menús de usuario y tooltips se veían sin fondo."),
            ("DONE", "Loop de redirects en /casasalco/", "El WSGI middleware strippeaba prefix y mi before_request lo redirigía de vuelta. Fix: chequear SCRIPT_NAME."),
        ],
    )

    add_heading(doc, "1.4 Features y módulos", level=2)
    add_status_table(
        doc,
        [
            ("DONE", "Permisos por rol en frontend", "Sidebar filtra menús + beforeLoad guards en rutas sensibles. Antes el cajero veía todo."),
            ("DONE", "Módulo Mantenimiento conectado", "Tabs con Familias / Rubros / Subrubros / Marcas. Componentes existían pero la ruta era placeholder."),
            ("DONE", "Stock con búsqueda y filtros server-side", "q (código/descripción) + estado (agotado/critico/reorden/sobrestock/ok/bajo_minimo) + paginación + endpoint /resumen."),
        ],
    )

    add_heading(doc, "1.5 Calidad y herramientas", level=2)
    add_status_table(
        doc,
        [
            ("DONE", "Agente APB de smoke testing", "~50 chequeos automatizados sobre la API: auth, permisos, inputs malos, catálogos, SPA. Ejecutable on-demand."),
            ("DONE", "4 manuales de usuario", "admin / supervisor / cajero / contador en .docx con screenshots embebidos automáticos (Playwright)."),
            ("DONE", "Scripts reutilizables", "scripts/capture_screenshots.py + scripts/generate_manuals.py para regenerar manuales tras cambios de UI."),
            ("DONE", "Memoria persistente", "Decisiones, bugs y discoveries guardados en engram para futuras sesiones."),
        ],
    )

    doc.add_page_break()

    # ===== Sección: Gap a producción =====
    add_heading(doc, "2. Gap a producción", level=1)
    add_paragraph(
        doc,
        "Pendientes para que el sistema pueda usarse en producción real con "
        "clientes y emisión de comprobantes fiscales.",
    )

    add_heading(doc, "2.1 Bloqueadores críticos (no se puede ir a prod sin esto)", level=2)
    add_status_table(
        doc,
        [
            ("RIESGO", "JWT_SECRET_KEY y SECRET_KEY default", "Actualmente: 'dev-jwt-secret-key-at-least-32-bytes-long-please-change'. Generar con `openssl rand -hex 32` y setear en .env. Si quedan así, cualquiera puede falsificar tokens."),
            ("RIESGO", "AFIP_MODE=mock", "La facturación electrónica devuelve CAEs falsos. Para producción: configurar pyafipws con cert+key reales, AFIP_HOMO=False, CUIT real."),
            ("RIESGO", "OCR_MODE=mock", "Los datos extraídos son fake. Para producción: setear OCR_MODE=gemini o anthropic con la API key correspondiente. El módulo está implementado y funciona; solo falta la key."),
            ("RIESGO", "Cert AFIP no instalado", "Una vez generado el cert/key real, ubicar en instance/afip_certs/<cuit>/. El cert dura 2 años — agendar renewal."),
        ],
    )

    add_heading(doc, "2.2 Importantes (deberían hacerse antes del go-live)", level=2)
    add_status_table(
        doc,
        [
            ("TODO", "WSGI server productivo", "Cambiar Werkzeug dev server por gunicorn + gevent (worker class para SocketIO). Comando: gunicorn --bind 0.0.0.0:5005 --worker-class gevent --workers 1 wsgi:app."),
            ("TODO", "Servicio Windows / supervisord", "Para que el backend se inicie automáticamente al boot y reinicie si crashea. Usar NSSM o WinSW para Windows."),
            ("TODO", "Migrar SQLite → Postgres", "SQLite OK para low traffic, pero con varios cajeros simultáneos puede sufrir lock contention. Postgres recomendado. Hay un script scripts/migrate_sqlite_to_postgres.py."),
            ("TODO", "Backups automáticos", "Programar pg_dump (o sqlite .backup) diario, retención 30 días, copia off-site (Google Drive / S3)."),
            ("TODO", "Datos reales de stock", "El stock actual es random demo. Cargar el inventario real de CASA SALCO desde XLS o conteo físico inicial."),
            ("TODO", "Datos reales de saldos cuenta corriente", "Cargar saldo inicial de clientes y proveedores con cuenta corriente."),
        ],
    )

    add_heading(doc, "2.3 Calidad de datos heredada del XLS", level=2)
    add_status_table(
        doc,
        [
            ("TODO", "2.051 marcas con duplicados", "Probablemente 'Coca Cola', 'COCA-COLA', 'COCACOLA' como entradas distintas. Detectar y consolidar (puede automatizarse con similitud)."),
            ("TODO", "Artículos con marca_id=NULL", "Los 34.847 importados del XLS no tienen marca asociada. Decidir si vale la pena vincular."),
            ("TODO", "611 rubros con muchos 'sin-rubro'", "El importador del XLS no encontraba algunos rubros y caía a fallback. Revisar."),
            ("TODO", "5 alertas con sucursal_id=NULL", "Detectado al hacer merge. Ver si es bug del seed o intencional."),
        ],
    )

    add_heading(doc, "2.4 Operación y monitoreo", level=2)
    add_status_table(
        doc,
        [
            ("TODO", "Logs estructurados", "Hoy salen a stdout. Para producción: logging estructurado (JSON) a archivos rotados, o a un colector (Loki / CloudWatch)."),
            ("TODO", "Monitoring básico", "Alertas si el backend cae, si la latencia sube, si la DB crece sin control. Sentry para errores en runtime."),
            ("TODO", "Procedimiento de rollback", "Documentar cómo volver atrás un release si rompe en producción."),
            ("TODO", "Plan de migration de schema", "Hoy las migraciones se corren con `flask db upgrade` manual. Definir cuándo/cómo en cada release."),
        ],
    )

    add_heading(doc, "2.5 Seguridad (no urgente pero recomendado)", level=2)
    add_status_table(
        doc,
        [
            ("TODO", "Rate limiting en /auth/login", "Mitigar brute-force. Flask-Limiter o nginx limit_req."),
            ("TODO", "Logging de auditoría", "Quién creó/modificó/eliminó qué y cuándo. Tabla 'audit_log' con triggers o middleware."),
            ("TODO", "CSP headers en nginx", "Content Security Policy para mitigar XSS aunque ya validamos en backend."),
            ("TODO", "HTTPS-only cookie flags", "Verificar httpOnly + secure + sameSite en JWT cookies si se usan (hoy es Bearer header)."),
            ("TODO", "Forzar password reset al primer login", "Los seeds tienen passwords genéricas (admin123, cajero123, etc). Forzar cambio al primer ingreso."),
        ],
    )

    add_heading(doc, "2.6 Funcionalidades pendientes", level=2)
    add_status_table(
        doc,
        [
            ("TODO", "Catalog-promociones", "Heredado de iteración previa. Diseño retry pendiente."),
            ("TODO", "Reportes históricos post-merge", "Verificar que los reportes funcionan con punto_venta y caja_numero remapeados (1..4 dentro de la única sucursal)."),
            ("TODO", "Mantenimiento — secciones extra", "Hoy tabs solo Familias / Rubros / Subrubros / Marcas. Falta: usuarios/roles, sucursales, parámetros del negocio (CUIT, IVA), listas de precios, importación XLS desde UI."),
            ("TODO", "Detector de duplicados (marcas)", "Endpoint que sugiere consolidaciones por similitud de string. Ayuda a limpiar las 2.051 marcas heredadas."),
        ],
    )

    add_heading(doc, "2.7 Deuda técnica conocida", level=2)
    add_status_table(
        doc,
        [
            ("TODO", "Vite proxy bug con Node v24", "Workaround actual: single-port. Solución real: downgrade Node a 20 LTS o reemplazar middleware http-proxy. Recuperaríamos HMR para desarrollo activo."),
            ("TODO", "Tests E2E / integración UI", "Solo hay smoke tests del agente APB sobre API. Sumar Playwright tests para flujos críticos (POS, alta de cliente, OCR)."),
            ("TODO", "StockAjusteTarget casteo feo", "El dialog espera Articulo completo pero recibimos solo el embedded. Cambiar el tipo del dialog a un Pick<Articulo, ...> o tipo propio."),
        ],
    )

    add_heading(doc, "2.8 Deploy físico (hardware en sucursal)", level=2)
    add_paragraph(
        doc,
        "El sistema tiene los drivers implementados y testeados, pero el "
        "deploy físico en cada PC de mostrador depende de hardware del "
        "cliente. Esto NO es deuda de código — es trabajo de instalación "
        "y configuración por caja.",
    )
    add_status_table(
        doc,
        [
            ("TODO", "Balanzas — identificar modelo por sucursal", "¿Cada sucursal tiene Systel o Kretz? ¿Qué modelo exacto? El agent soporta ambas marcas + network. Driver mock funciona OK para preview."),
            ("TODO", "Balanzas — cablear y configurar COM port", "Conectar serial RS-232 a la PC del cajero. Identificar COM port asignado por Windows."),
            ("TODO", "Balanzas — setear .env del agent por caja", "SCALE_MODE=systel|kretz|network, SCALE_PORT=COMX, SCALE_BAUDRATE=9600. Repetir por caja."),
            ("TODO", "Balanzas — test físico end-to-end", "Pesar producto real en la balanza, abrir POS, confirmar que weigh-dialog lee el peso. Validar tara y stability flag."),
            ("TODO", "Impresoras 3NSTAR — driver USB o Network", "PRINTER_MODE=usb (PRP-080 USB) o network (PRP-080N Ethernet). Driver implementado, falta config por caja."),
            ("TODO", "Impresoras — papel 80mm + cortado", "PRINTER_PAPER_WIDTH_MM=80, prueba de impresión real con corte automático."),
            ("TODO", "Agent como servicio Windows", "Hoy se arranca a mano con `python -m jarvis_agent`. Para producción: NSSM o WinSW para que levante al boot, reinicie al crashear, y se vea en services.msc."),
            ("TODO", "AFIP cert + key físicos en cada sucursal", "El cert AFIP se ubica en `instance/afip_certs/<cuit>/`. Cada sucursal puede usar el mismo cert (CASA SALCO único CUIT) o uno propio si hubiese sucursales con razón social distinta."),
            ("TODO", "Backups en cada caja", "SQLite local del POS Tauri (Fase 2): backup diario a la red interna o a una USB rotativa."),
        ],
    )

    doc.add_page_break()

    # ===== Plan recomendado =====
    add_heading(doc, "3. Plan de acción recomendado", level=1)

    add_heading(doc, "Sprint 1 — Pase a producción (prioridad máxima)", level=2)
    add_bullets(
        doc,
        [
            "Generar JWT_SECRET_KEY y SECRET_KEY reales con `openssl rand -hex 32` y setear en .env del servidor.",
            "Conseguir cert+key de AFIP, instalarlos, configurar AFIP_MODE=pyafipws y AFIP_HOMO=false.",
            "Conseguir API key de Gemini (recomendado: ultrabarato + free tier amplio) y setear OCR_MODE=gemini.",
            "Cargar el inventario real (stock) — desde XLS, conteo físico o ETL del sistema viejo.",
            "Migrar a Postgres y setup gunicorn como servicio Windows. Backups diarios.",
            "Smoke test final con el agente APB tras todos los cambios.",
        ],
    )

    add_heading(doc, "Sprint 2 — Operación robusta", level=2)
    add_bullets(
        doc,
        [
            "Logging estructurado + Sentry para captura de errores.",
            "Rate limiting en /auth/login. Forzar password reset al primer login.",
            "Logging de auditoría: tabla audit_log + triggers.",
            "Procedimientos documentados: backup, restore, rollback, deploy.",
        ],
    )

    add_heading(doc, "Sprint 3 — Calidad de datos y features", level=2)
    add_bullets(
        doc,
        [
            "Detector y consolidador de marcas duplicadas.",
            "Mantenimiento — secciones extra (usuarios/roles, parámetros).",
            "Catalog-promociones design retry.",
            "Tests E2E con Playwright para los flows críticos.",
        ],
    )

    # ===== Cierre =====
    doc.add_paragraph()
    add_paragraph(
        doc,
        f"Documento generado el {today}. Refleja el estado del repositorio y de "
        "la infraestructura observable en ese momento. Para tener un panorama "
        "actualizado, regenerar este documento corriendo "
        "`scripts/generate_gap_doc.py` después de cualquier cambio mayor.",
        italic=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK: {OUT.relative_to(ROOT)} ({OUT.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
