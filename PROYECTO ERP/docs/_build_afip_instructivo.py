"""
Genera un instructivo Word paso a paso para obtener Cert + Key AFIP
y configurarlos en Jarvis Core.

Uso:
    python _build_afip_instructivo.py
Salida:
    Jarvis-Core-Instructivo-AFIP.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "..", "Jarvis-Core-Instructivo-AFIP.docx"
)
OUT = os.path.abspath(OUT)

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.4)
    s.right_margin = Cm(2.4)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


# ================ helpers ================
def H1(text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(10)
    return h

def H2(text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    return h

def H3(text):
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h

def P(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p

def B(text):
    """Bullet."""
    return doc.add_paragraph(text, style='List Bullet')

def N(text):
    """Numbered list item."""
    return doc.add_paragraph(text, style='List Number')

def CODE(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    return p

def CALLOUT(title, text, color="warning"):
    """Color: warning (amber), info (blue), critical (red), ok (green)."""
    color_map = {
        "warning": (0xFE, 0xF3, 0xC7),
        "info": (0xDB, 0xEA, 0xFE),
        "critical": (0xFE, 0xE2, 0xE2),
        "ok": (0xD1, 0xFA, 0xE5),
    }
    bg = color_map.get(color, color_map["info"])
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    # Set background
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*bg))
    tc_pr.append(shd)
    p = cell.paragraphs[0]
    r = p.add_run(title + ":")
    r.bold = True
    p.add_run(" " + text)
    doc.add_paragraph("")
    return t

def TABLE(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = 'Light Grid Accent 1'
    except KeyError:
        t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        rr = p.add_run(h)
        rr.bold = True
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = str(v) if v is not None else ""
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph("")
    return t

def BREAK():
    doc.add_page_break()


# ================ COVER ================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("\n\n\n\nJARVIS CORE\n")
tr.bold = True
tr.font.size = Pt(36)
tr.font.color.rgb = RGBColor(0x00, 0x7A, 0xFF)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Instructivo AFIP")
sr.font.size = Pt(20)
sr.italic = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr2 = sub2.add_run(
    "Como obtener Certificado Digital y Clave Privada para "
    "facturacion electronica en produccion"
)
sr2.font.size = Pt(13)

doc.add_paragraph("\n" * 6)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Para: Laura Fabiana Castro - CUIT 27-22515199-2")
mr.font.size = Pt(12)
mr.bold = True

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
nr = note.add_run("Tiempo estimado: 1 a 2 horas")
nr.font.size = Pt(11)
nr.italic = True

BREAK()


# ================ INDICE ================
H1("Indice")
INDEX = [
    "1. Para que sirve este tramite",
    "2. Requisitos previos (que necesitas tener antes de arrancar)",
    "3. Conceptos claves (que es cada cosa)",
    "4. Paso 1 - Instalar OpenSSL",
    "5. Paso 2 - Generar clave privada y solicitud (CSR)",
    "6. Paso 3 - Adherir el servicio WSFE en AFIP",
    "7. Paso 4 - Subir el CSR y descargar el certificado",
    "8. Paso 5 - Asociar el certificado al servicio WSFE",
    "9. Paso 6 - Habilitar punto de venta",
    "10. Paso 7 - Configurar Jarvis Core con tu certificado",
    "11. Paso 8 - Probar emision real",
    "12. Paso 9 - Homologacion (los 7 casos de prueba)",
    "13. Renovacion del certificado (cada 2 anos)",
    "14. Problemas frecuentes y soluciones",
    "15. Recursos y contactos",
]
for line in INDEX:
    p = doc.add_paragraph()
    p.add_run(line).font.size = Pt(11)

BREAK()


# ================ 1. Para que sirve ================
H1("1. Para que sirve este tramite")

P(
    "Hoy Jarvis Core esta funcionando con CAEs simulados (modo \"mock\"). Los "
    "comprobantes que emite tienen un CAE inventado para que vos puedas probar "
    "el sistema. Esto NO es legal para vender ante AFIP."
)
P(
    "Para emitir comprobantes con validez fiscal real necesitas que AFIP te "
    "entregue un Certificado Digital y una Clave Privada. Con esos dos "
    "archivos, Jarvis Core puede pedirle CAEs reales a AFIP en cada venta."
)

H3("Que vas a tener al final de este instructivo")
B("Un archivo certificado.crt con tu identidad digital ante AFIP.")
B("Un archivo privada.key con tu clave secreta (solo vos la tenes).")
B("Tu Punto de Venta habilitado (numero, ej. 0001 para tu sucursal principal).")
B("Jarvis Core configurado para emitir CAEs reales.")
B("Comprobantes legales con QR AFIP impresos en la comandera 3NSTAR.")

CALLOUT(
    "Importante",
    "Los archivos privada.key y certificado.crt son ULTRA SENSIBLES. Cualquiera "
    "que los tenga puede emitir facturas a tu nombre. Guardalos en un lugar "
    "seguro y NUNCA los compartas por mail/WhatsApp.",
    color="critical"
)


# ================ 2. Requisitos previos ================
H1("2. Requisitos previos")

P("Antes de arrancar, asegurate de tener:")

TABLE(
    ["Requisito", "Como conseguirlo"],
    [
        ["CUIT activo", "Ya lo tenes: 27-22515199-2 (Laura Fabiana Castro)"],
        ["Clave Fiscal nivel 3", "Si no la tenes, sacar turno en cualquier dependencia AFIP. Sin esto NO podes seguir."],
        ["PC con acceso a internet", "Cualquier Windows / Mac / Linux"],
        ["OpenSSL instalado", "Lo vamos a instalar en el Paso 1"],
        ["Tiempo libre (1 a 2 horas)", "Hace el tramite en una sentada para no perder el hilo"],
        ["Acceso a backend/.env de Jarvis Core", "Lo configuramos al final"],
    ],
    widths=[5.5, 9],
)

CALLOUT(
    "Antes de seguir",
    "Verifica que tu Clave Fiscal sea nivel 3 entrando a https://auth.afip.gob.ar. "
    "En \"Mi Perfil\" deberia decir \"Nivel de seguridad: 3\". Si dice 2 o menor, "
    "sacar turno en AFIP para upgrade.",
    color="warning"
)


# ================ 3. Conceptos claves ================
H1("3. Conceptos claves")

TABLE(
    ["Termino", "Que significa en plata"],
    [
        ["Certificado Digital (.crt)", "Tu DNI digital ante AFIP. Lo emite AFIP basandose en una solicitud que vos firmas con tu clave privada."],
        ["Clave Privada (.key)", "La contrasena que solo vos tenes. Se usa para firmar pedidos a AFIP. NUNCA se comparte."],
        ["CSR (Certificate Signing Request)", "El \"formulario\" firmado con tu clave privada que le mandas a AFIP para que te genere el certificado."],
        ["WSAA", "Web Service de Autenticacion AFIP. Es el primer servicio que llamas: te entrega un \"Ticket de Acceso\" (TA) valido por 12 horas."],
        ["WSFE / WSFEv1", "Web Service de Factura Electronica. Es el servicio que entrega el CAE para cada comprobante."],
        ["TA (Ticket de Acceso)", "Es como una sesion. Jarvis Core lo pide al WSAA y lo cachea por 12 horas. Vos no tenes que hacer nada con esto."],
        ["CAE (Codigo de Autorizacion Electronico)", "El numero que AFIP entrega para autorizar UN comprobante. Cada factura B/C que emitas va a tener su CAE."],
        ["Punto de Venta", "Numero (ej. 0001, 0002) que identifica una caja/terminal/sucursal ante AFIP. Cada uno tiene su numeracion contigua propia."],
        ["Homologacion", "Ambiente de PRUEBAS de AFIP. Los CAEs son reales pero los comprobantes no tienen validez fiscal. Ideal para probar."],
        ["Produccion", "Ambiente REAL. Los CAEs son legales y los comprobantes valen para la AFIP."],
    ],
    widths=[5, 9.5],
)

BREAK()


# ================ Paso 1 ================
H1("4. Paso 1 - Instalar OpenSSL")

P(
    "OpenSSL es la herramienta que vamos a usar para generar la clave privada y "
    "el CSR. Es estandar y gratis."
)

H3("Windows")
N("Ir a https://slproweb.com/products/Win32OpenSSL.html")
N("Descargar la version \"Win64 OpenSSL v3.x.x\" (la mas nueva, NO la \"Light\")")
N("Ejecutar el instalador")
N("Dejar todas las opciones por default. Cuando pregunte \"Copy OpenSSL DLLs to:\", elegi \"The OpenSSL binaries (/bin) directory\"")
N("Reiniciar la PowerShell o terminal para que tome el PATH")
N("Verificar abriendo PowerShell y tipeando: openssl version")
P("Deberia salir algo como:")
CODE("OpenSSL 3.x.x  XX Mmm YYYY (Library: OpenSSL 3.x.x XX Mmm YYYY)")

H3("Mac")
CODE("brew install openssl")

H3("Linux (Debian/Ubuntu)")
CODE("sudo apt install openssl")

CALLOUT(
    "Tip",
    "Si openssl version te tira error \"command not found\", el instalador no agrego "
    "el path. En Windows: Sistema -> Variables de entorno -> Path -> Agregar "
    "C:\\Program Files\\OpenSSL-Win64\\bin -> Aceptar -> reiniciar terminal.",
    color="info"
)


# ================ Paso 2 ================
H1("5. Paso 2 - Generar clave privada y CSR")

P(
    "Vamos a generar dos archivos: privada.key (tu clave) y solicitud.csr "
    "(la solicitud que vamos a subir a AFIP)."
)

H3("Crear una carpeta de trabajo")
P("Crea una carpeta nueva, por ejemplo:")
CODE("C:\\afip-jarvis\\")
P("Abri PowerShell y entra a esa carpeta:")
CODE("cd C:\\afip-jarvis")

H3("Generar la clave privada (1024 bytes, formato PEM)")
CODE("openssl genrsa -out privada.key 2048")
P("Resultado: archivo privada.key en la carpeta. NUNCA borres este archivo.")

H3("Generar el CSR (la solicitud para AFIP)")
P("Comando con los datos de tu CUIT (todo en una sola linea):")
CODE(
    'openssl req -new -key privada.key '
    '-subj "/C=AR/O=Laura Fabiana Castro/CN=jarvis-core-prod/serialNumber=CUIT 27225151992" '
    '-out solicitud.csr'
)
P("Detalle de los campos:")
B("/C=AR -> Pais: Argentina")
B("/O=Laura Fabiana Castro -> Razon social. PEGALO TAL CUAL ESTA INSCRIPTO EN AFIP.")
B("/CN=jarvis-core-prod -> Alias del certificado. Este nombre es solo para tu referencia.")
B("/serialNumber=CUIT 27225151992 -> El CUIT SIN guiones. Critico que este bien.")

P("Ahora deberias tener en C:\\afip-jarvis\\:")
CODE(
    "privada.key       (NO COMPARTIR — guardar en lugar seguro)\n"
    "solicitud.csr     (esto es lo que vas a subir a AFIP)"
)

CALLOUT(
    "Backup IMPORTANTE",
    "Copia privada.key a una segunda ubicacion segura (USB, Drive privado, etc.) "
    "ANTES de seguir. Si perdes este archivo, tenes que reiniciar todo el tramite "
    "desde cero.",
    color="critical"
)


# ================ Paso 3 ================
H1("6. Paso 3 - Adherir el servicio WSFE en AFIP")

P(
    "Antes de generar el certificado, hay que decirle a AFIP que vas a usar el "
    "servicio de Factura Electronica."
)

N("Ir a https://auth.afip.gob.ar/")
N("Login con tu CUIT 27225151992 + clave fiscal")
N("Buscar y abrir el servicio \"Administrador de Relaciones de Clave Fiscal\". Si no lo ves en la lista, hay que adherirlo primero clickeando \"Adherir Servicio\" -> AFIP -> \"Administrador de Relaciones\"")
N("Una vez dentro: click en boton \"Nueva Relacion\"")
N("\"Servicio\": click la lupa -> seleccionar AFIP -> \"WebServices\" -> \"WSFE - Facturacion Electronica\"")
N("\"Representado\": deja el default (vos misma)")
N("\"Computador Fiscal\": dejar vacio por ahora (lo asociamos al cert despues)")
N("Confirmar")

CALLOUT(
    "Si te aparece error",
    "Si sale \"el servicio ya esta adherido\", esta perfecto. Pasamos al siguiente paso.",
    color="info"
)


# ================ Paso 4 ================
H1("7. Paso 4 - Subir el CSR y descargar el certificado")

P(
    "Ahora vamos a subir tu solicitud.csr a AFIP y obtener el certificado.crt."
)

N("Volver a https://auth.afip.gob.ar/ (si te deslogueo, login de nuevo)")
N("Buscar el servicio \"Administracion de Certificados Digitales\". Si no esta, adherirlo igual que el WSFE: \"Adherir Servicio\" -> AFIP -> \"Administracion de Certificados Digitales\"")
N("Una vez dentro: click \"Agregar alias\"")
N("\"Alias\": jarvis-core-prod (el mismo que pusiste en el CN del CSR)")
N("\"Solicitud (CSR)\": seleccionar el archivo C:\\afip-jarvis\\solicitud.csr")
N("Confirmar")
N("AFIP va a generar el certificado en el momento. Aparece en la lista.")
N("Click en el boton \"Ver\" del certificado recien creado")
N("Hacer click derecho -> \"Guardar como\" sobre el contenido")
N("Guardar como certificado.crt en C:\\afip-jarvis\\")

P("Ahora deberias tener en C:\\afip-jarvis\\:")
CODE(
    "privada.key\n"
    "solicitud.csr\n"
    "certificado.crt   (NUEVO — el que te dio AFIP)"
)


# ================ Paso 5 ================
H1("8. Paso 5 - Asociar el certificado al servicio WSFE")

P(
    "Ya tenes el certificado, pero AFIP necesita saber que ese certificado puede "
    "usar el servicio WSFE."
)

N("Volver a \"Administrador de Relaciones de Clave Fiscal\"")
N("Click en \"Nueva Relacion\"")
N("\"Servicio\": elegir AFIP -> WebServices -> WSFE - Facturacion Electronica")
N("\"Representado\": tu CUIT")
N("\"Computador Fiscal\": click la lupa -> seleccionar tu certificado (jarvis-core-prod)")
N("Confirmar")

CALLOUT(
    "Verificacion",
    "Despues de confirmar, en la lista de relaciones tendrias que ver una entrada "
    "que dice algo asi: WSFE -> jarvis-core-prod -> 27225151992. Si la ves, anda "
    "perfecto.",
    color="ok"
)


# ================ Paso 6 ================
H1("9. Paso 6 - Habilitar punto de venta")

P(
    "Cada terminal/sucursal donde emitis facturas necesita un \"punto de venta\" "
    "registrado ante AFIP. Para arrancar con UNA sucursal, registramos el punto 0001."
)

H3("En el portal AFIP")
N("Volver a https://auth.afip.gob.ar/")
N("Buscar \"Administracion de puntos de venta y domicilios\"")
N("\"Agregar\" un punto de venta")
N("\"Numero\": 1 (queda como 00001)")
N("\"Nombre de Fantasia\": el de tu sucursal principal (ej. \"Castulo Centro\")")
N("\"Sistema\": elegir \"RECE para aplicativo y web services\"")
N("\"Domicilio\": agregar el de la sucursal")
N("Confirmar")

CALLOUT(
    "Si tenes 4 sucursales",
    "Vas a necesitar agregar 4 puntos de venta (00001, 00002, 00003, 00004), uno "
    "por cada una. Pero podes empezar con uno solo y agregar los demas mas adelante. "
    "El numero de punto de venta es ETERNO — una vez asignado a una sucursal, NO se puede "
    "reusar para otra en el mismo periodo fiscal.",
    color="warning"
)


# ================ Paso 7 ================
H1("10. Paso 7 - Configurar Jarvis Core con tu certificado")

P(
    "Ya tenes los archivos. Ahora le decimos a Jarvis Core donde estan."
)

H3("Copiar los archivos al servidor")

P("Crear la carpeta:")
CODE("D:\\repo\\00-omar\\Castulo\\nuevo\\backend\\instance\\afip_certs\\27225151992\\")

P("Copiar los archivos:")
CODE(
    "privada.key       -> instance\\afip_certs\\27225151992\\privada.key\n"
    "certificado.crt   -> instance\\afip_certs\\27225151992\\certificado.crt"
)

CALLOUT(
    "Permisos",
    "En Linux/Mac, hace 'chmod 600' a privada.key. En Windows asegurate que solo "
    "el usuario que corre Flask pueda leer la carpeta.",
    color="warning"
)

H3("Editar backend/.env")

P("Cambiar (o agregar) estas lineas:")
CODE(
    "AFIP_MODE=pyafipws\n"
    "AFIP_HOMO=false\n"
    "AFIP_CUIT=27225151992\n"
    "AFIP_CERT_PATH=instance/afip_certs/27225151992/certificado.crt\n"
    "AFIP_KEY_PATH=instance/afip_certs/27225151992/privada.key"
)

CALLOUT(
    "Importante - homologacion",
    "Recomiendo arrancar con AFIP_HOMO=true (ambiente de pruebas) hasta que "
    "completes el Paso 9 (homologacion con los 7 casos de prueba). Solo cambia a "
    "false cuando AFIP te valide el resultado.",
    color="warning"
)

H3("Reiniciar el backend")

CODE(
    "cd D:\\repo\\00-omar\\Castulo\\nuevo\\backend\n"
    ".venv\\Scripts\\python.exe wsgi.py"
)


# ================ Paso 8 ================
H1("11. Paso 8 - Probar emision real")

P("Hace una venta de prueba para validar que todo esta conectado:")

N("Login en http://localhost:5173 con admin@castulo.app")
N("Ir a Punto de Venta")
N("Seleccionar la sucursal SUC01")
N("Cargar 1-2 articulos al carrito")
N("Tipo de comprobante: Factura B")
N("Cliente: Consumidor Final")
N("Pago: efectivo por el total con IVA")
N("F12 para finalizar")
N("En el modal post-venta: click \"Imprimir ticket\"")

P("Que mirar:")
B("El modal deberia mostrar un CAE de 14 digitos (NO el mock que era hash sha256)")
B("La fecha de vencimiento del CAE deberia ser 10 dias en el futuro")
B("El PDF impreso deberia tener QR AFIP funcional (escanealo con el celular -> tiene que abrir la pagina de verificacion AFIP)")

CALLOUT(
    "Si algo falla",
    "Ve la seccion 14 \"Problemas frecuentes\". El error mas comun es timezone "
    "del PC desincronizado.",
    color="info"
)


# ================ Paso 9 ================
H1("12. Paso 9 - Homologacion (los 7 casos de prueba)")

P(
    "Antes de pasar a produccion, AFIP exige que valides el sistema con 7 casos "
    "de prueba en el ambiente de homologacion. Esto se llama \"homologacion\"."
)

P("Si AFIP_HOMO=true, los CAEs que emitis van al ambiente de prueba. Sirve para validar.")

H3("Los 7 casos")
TABLE(
    ["#", "Tipo", "Detalle"],
    [
        ["1", "Factura A", "RI a RI - importe normal"],
        ["2", "Factura A con percepcion", "Con percepcion IIBB"],
        ["3", "Factura B", "Consumidor Final - importe bajo"],
        ["4", "Factura B importe alto", "Mayor a $417.119 (Oct 2024) - obliga CondicionIVAReceptorId"],
        ["5", "Nota de Credito A", "Sobre la factura del caso 1"],
        ["6", "Nota de Debito A", "Sobre la factura del caso 1"],
        ["7", "Factura C", "Solo si emitis a Monotributistas"],
    ],
    widths=[1, 4, 9.5],
)

P("Como hacerlo:")
N("Asegurate AFIP_HOMO=true en backend/.env")
N("Ir al portal AFIP, descargar la planilla \"Casos de Homologacion\" del manual WSFE")
N("Emitir los 7 casos desde Jarvis Core, anotar los CAEs que devuelve")
N("Subir la planilla con los CAEs al portal AFIP, en \"Solicitar Habilitacion\"")
N("Esperar 24-72 horas hasta que AFIP confirme")
N("Una vez aprobado, cambiar AFIP_HOMO=false en .env")
N("Reiniciar backend")
N("Listo - estas en produccion REAL")


# ================ Renovacion ================
H1("13. Renovacion del certificado (cada 2 anos)")

P(
    "El certificado dura exactamente 2 anos desde el dia que lo emitiste. AFIP "
    "permite renovarlo a partir de los 60 dias previos al vencimiento."
)

H3("Recordatorios sugeridos")
B("Agendar en tu calendario (Google Calendar / Outlook): \"AFIP cert vence en 60 dias\" -> 60 dias antes del vencimiento real")
B("Y otro: \"AFIP cert vence MANANA\" -> 1 dia antes del vencimiento")

H3("Proceso de renovacion")
N("Generar nueva clave privada y CSR (igual que en el Paso 2, podes usar la misma carpeta)")
N("En el portal AFIP: \"Administracion de Certificados Digitales\" -> click \"Renovar\" sobre el alias")
N("Subir el nuevo CSR")
N("Descargar el nuevo certificado")
N("Hacer la asociacion con WSFE como en el Paso 5")
N("Reemplazar los archivos en instance/afip_certs/27225151992/")
N("Reiniciar backend")

CALLOUT(
    "Backup",
    "Mantene el certificado VIEJO en una subcarpeta como instance/afip_certs/27225151992/old/. "
    "Si algo falla con el nuevo, podes volver al viejo (siempre que no haya vencido).",
    color="info"
)


# ================ Problemas frecuentes ================
H1("14. Problemas frecuentes y soluciones")

H3("Error: \"Invalid TA generationTime\"")
B("Causa: El reloj de tu PC esta desincronizado. AFIP rechaza tickets con generationTime fuera de rango (5 minutos de tolerancia).")
B("Solucion Windows: Configuracion -> Hora e idioma -> Sincronizar ahora. Activar \"Establecer la hora automaticamente\".")
B("Solucion Linux: sudo timedatectl set-ntp true")

H3("Error: \"CUIT representado no autorizado\"")
B("Causa: El servicio WSFE no esta adherido al CUIT, o no esta asociado al certificado.")
B("Solucion: revisar Pasos 3 y 5.")

H3("Error: \"Punto de venta no autorizado\"")
B("Causa: El punto de venta que mandas en la solicitud no esta habilitado en AFIP.")
B("Solucion: revisar Paso 6. En Jarvis Core, asegurate que la sucursal usa el numero de punto de venta correcto (en el modulo Sucursales).")

H3("Error: \"Numero de comprobante invalido\"")
B("Causa: El numero que mandaste no es contiguo con el ultimo autorizado.")
B("Solucion: Jarvis Core usa FECompUltimoAutorizado al iniciar. Reiniciar el backend para que vuelva a sincronizar la numeracion con AFIP.")

H3("Error: \"Token o sign invalidos\"")
B("Causa: La firma del TA fallo. Generalmente clave privada incorrecta o cert vencido.")
B("Solucion: chequear que privada.key y certificado.crt estan en la carpeta correcta y son los mismos archivos del par (no mezclar de tramites distintos).")

H3("Error: \"Servicio no disponible\"")
B("Causa: AFIP esta caido (suele pasar cuando vienen muchas facturas a fin de mes o a fin de ejercicio).")
B("Solucion: esperar 5-30 minutos. Jarvis Core va a hacer retry automatico.")

H3("Error: \"CondicionIVAReceptorId obligatorio\"")
B("Causa: Desde Octubre 2024 (RG 5616) AFIP obliga este campo en facturas B mayores a $417.119.")
B("Solucion: en el cliente, completar el campo \"Condicion IVA Receptor\". Para Consumidor Final = 5.")


# ================ Recursos ================
H1("15. Recursos y contactos")

H3("Documentacion oficial")
B("Portal AFIP: https://www.afip.gob.ar")
B("Manual desarrollador WSFE: https://www.afip.gob.ar/fe/documentos/manual_desarrollador_COMPG_v2_10.pdf")
B("Tablas de parametros AFIP: https://www.afip.gob.ar/fe/documentos/TablasParametros.docx")

H3("Comunidad y soporte")
B("PyAfipWs (libreria): https://github.com/PyAfipWs/pyafipws")
B("Foro Google de PyAfipWs: https://groups.google.com/g/pyafipws")
B("Centro de ayuda AFIP: 0810-999-2347")

H3("Si te trabas")
P(
    "Si en algun paso te trabas, lo mejor es contactar a un Contador con "
    "experiencia en webservices AFIP. El tramite tecnico no es complicado pero "
    "tiene burocracia AFIP que un contador resuelve en minutos."
)


# ================ FIN ================
doc.add_paragraph("")
fin = doc.add_paragraph()
fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fin.add_run("--- FIN DEL INSTRUCTIVO ---")
fr.bold = True
fr.font.size = Pt(12)

doc.add_paragraph("")
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.add_run("Generado por Jarvis Core - v0.1.0").italic = True


doc.save(OUT)
print(f"OK -> {OUT}")
print(f"Tamano: {os.path.getsize(OUT) / 1024:.1f} KB")
